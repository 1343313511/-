#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""仅填充实习主要内容和实习总结，其他内容原封不动"""

from docx import Document
from docx.shared import Pt

template_path = '/root/.openclaw/media/qqbot/downloads/1905229424/17911036127897E831347AF9570D93DF/0bfb8723-348a-4af9-898c-71e71ea272d8.docx'
output_path = '/opt/Class01/项目/网络安全实践-2实习报告-王宇杨.docx'

doc = Document(template_path)
table = doc.tables[0]
cell = table.rows[0].cells[0]

# ============ 实习主要内容（P7-P25, 共19个空段落） ============
main_lines = [
    "本次实训以 Kali Linux 为操作平台，以一个完整的 Flask Web 用户管理系统为靶标项目，从攻防双视角系统性地开展了为期九天的安全实训。实训覆盖了 OWASP Top 10 中绝大部分核心漏洞类型及渗透测试进阶技术。具体内容如下：",
    "",
    "1. 环境搭建与基础配置（Day 1）：完成 Kali Linux 系统初始化配置与工具链安装；部署 OpenClaw AI 系统并对接 QQ Bot 通道，实现通过 QQ 与 AI Agent 交互的能力；将项目 Fork 至个人 GitHub 仓库，建立 Git 版本管理规范。",
    "",
    "2. 密码安全与暴力破解防御（Day 1-2）：梳理常见系统默认账号密码，建立安全意识基线；配置 Burp Suite 代理并掌握请求拦截、修改、重放等核心功能；使用 Burp Suite Intruder 模块搭配字典进行暴力破解攻击实验。防护措施：在项目中实现指数退避延迟算法、图形验证码、账号锁定机制（5次失败锁定15分钟）、密码哈希存储及时间线攻击防护。",
    "",
    "3. SQL 注入漏洞修复（Day 2）：项目初始版本使用 f-string 拼接 SQL 语句，存在严重注入风险。攻击者可输入 admin 万能密码绕过认证或在搜索接口使用 UNION SELECT 读取任意表数据。修复方案：将所有 SQL 操作改为参数化查询，明文密码改为哈希存储，限制用户输入长度。",
    "",
    "4. 越权漏洞修复（Day 2）：存在垂直越权漏洞，未登录用户可直接访问 /profile、/upload 等需认证页面。修复方案：新增 login_required 装饰器，未登录自动重定向至登录页；同时增加水平越权防护，校验当前用户仅能访问自己的数据。",
    "",
    "5. CSRF 漏洞修复（Day 3）：所有需认证的 POST 操作存在 CSRF 风险。修复方案：使用 Flask-WTF 的 CSRFProtect 中间件全局启用 CSRF 令牌验证，所有表单添加隐藏的 csrf_token 字段。",
    "",
    "6. 文件上传漏洞修复（Day 2-3）：存在路径穿越和仅检查 Content-Type 头两种绕过方式。修复方案：使用 werkzeug.secure_filename() 处理文件名，新增文件魔数（Magic Number）校验验证真实文件类型，限制上传文件大小为 2MB。",
    "",
    "7. 文件包含与 XSS 漏洞修复（Day 4）：/page 路由存在本地文件包含漏洞，昵称、搜索框等多处存在 XSS 漏洞。修复方案：文件包含使用 os.path.normpath + os.path.realpath 双重路径校验；XSS 修复使用 html.escape() 对所有用户输入进行 HTML 实体转义。",
    "",
    "8. SSRF 漏洞修复（Day 5）：/fetch-url 路由服务端未对用户提交 URL 做任何过滤，可访问内网服务。修复方案：URL 协议白名单（仅 http/https）、内网地址黑名单过滤、自定义 DNS 解析验证、socket 超时控制。",
    "",
    "9. SSTI 模板注入修复（Day 6）：/welcome 和 /feedback 使用 render_template_string 拼接用户输入，可注入模板语法泄露密钥或通过对象继承链执行任意代码实现 RCE。修复方案：新增 safe_str() 函数对所有用户输入执行 html.escape() 转义后再拼接。",
    "",
    "10. 命令注入修复（Day 7）：/ping 路由使用 f-string 拼接命令并设置 shell=True，攻击者可注入 shell 元字符执行任意命令。修复方案：白名单字符正则校验、危险关键词拦截、改为参数列表形式彻底禁用 shell=True。",
    "",
    "11. Shell 反弹与 Linux 利用手法：掌握 bash / nc / Python / PHP 等多种反向 Shell 的构造方式；编写 PHP 一句话木马结合反向 Shell 实现 WebShell 管理；学习 SUID 提权、计划任务利用、sudo 配置错误等 Linux 系统层面攻击手法；配置 NPS 内网穿透客户端实现外网访问内网服务；针对 SQLi-LABS 靶场开展 WAF 绕过模糊测试。",
]

for i, text in enumerate(main_lines):
    idx = 7 + i
    p = cell.paragraphs[idx]
    for run in list(p.runs):
        run.text = ''
    if text:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'

# ============ 实习总结（P27-P63, 共37个空段落） ============
summary_lines = [
    "通过本次为期九天的网络安全实训，我系统性地构建了 Web 安全的知识体系，从最初仅了解漏洞概念，到能够独立分析漏洞原理、编写攻击载荷、完成修复方案并输出专业漏洞报告，全过程覆盖了 SQL 注入、越权、CSRF、文件上传绕过、文件包含、XSS、SSRF、SSTI、命令注入等 OWASP Top 10 核心漏洞类型。",
    "在工具使用方面，熟练掌握了 Kali Linux 操作系统环境、Burp Suite 进行 HTTP 流量分析与字典攻击、OpenClaw AI Agent 辅助漏洞分析与报告生成、Git 进行版本管理与协同开发等安全从业人员必备技能。",
    "本次实训的一大特色是将 AI Agent 深度集成到安全工作流中。AI 不仅辅助完成了漏洞分析、修复代码编写和报告文档的自动生成，还通过自动化工具有条理地管理项目文件、执行 Git 操作，极大提升了工作效率。这让我深刻认识到 AI 在安全领域的应用前景——AI 不会替代安全工程师，但能大幅提升安全工作的效率和质量。",
    "实训最大的收获是安全思维的建立。Web 应用的安全性由每一个输入点、每一个输出点、每一条数据库查询、每一次系统调用共同决定。安全不是靠单个功能实现的，它需要渗透到开发的全生命周期。九天的修复历程展示了安全是一个持续的、迭代的过程，这种持续安全的理念将贯穿我未来的职业生涯。",
    "本次实训共完成 20 次 Git 提交，生成 8 份独立漏洞修复报告（Word 格式），累计修复 9 大类安全漏洞，实现了从初始版本到全面安全加固的完整演进。项目代码及全部文档已推送至 GitHub 仓库存档。",
]

for i, text in enumerate(summary_lines):
    idx = 27 + i
    p = cell.paragraphs[idx]
    for run in list(p.runs):
        run.text = ''
    if text:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'

doc.save(output_path)
print(f'✅ 完成: {output_path}')
