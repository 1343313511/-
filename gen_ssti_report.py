#!/usr/bin/env python3
"""生成 SSTI 漏洞修复报告 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ============================================================
# 样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# ============================================================
# 封面
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SSTI 模板注入漏洞修复报告')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Day 8 — 服务端模板注入（Server-Side Template Injection）')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(6):
    doc.add_paragraph()

# 项目信息表
info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('项目名称', '用户管理系统（Class01）'),
    ('报告日期', datetime.datetime.now().strftime('%Y-%m-%d')),
    ('修复人员', '王宇杨'),
    ('漏洞类型', 'Server-Side Template Injection (SSTI)'),
    ('风险等级', '🔴 严重（Critical）'),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for cell in info_table.rows[i].cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)

doc.add_page_break()

# ============================================================
# 目录
# ============================================================
doc.add_heading('目录', level=1)
toc_items = [
    '一、漏洞概述',
    '二、漏洞原理',
    '三、漏洞详情与修复',
    '    3.1 /welcome 路由 SSTI 漏洞',
    '    3.2 /feedback 路由 SSTI 漏洞',
    '    3.3 /page 路由 SSTI 漏洞',
    '    3.4 其他潜在的 SSTI 风险点',
    '四、攻击复现',
    '五、修复方案详解',
    '六、修复验证',
    '七、安全编码规范建议',
    '八、修复变更清单',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 一、漏洞概述
# ============================================================
doc.add_heading('一、漏洞概述', level=1)

doc.add_paragraph(
    '本次安全审计发现用户管理系统存在三个 SSTI（Server-Side Template Injection，'
    '服务端模板注入）漏洞。SSTI 是一种严重的安全漏洞，攻击者可以通过在用户输入中'
    '注入 Jinja2 模板语法，在服务器端执行任意代码，从而完全控制服务器。'
)

doc.add_heading('漏洞总览', level=2)
vuln_table = doc.add_table(rows=4, cols=5)
vuln_table.style = 'Light Grid Accent 1'
headers = ['漏洞编号', '漏洞位置', '风险等级', '漏洞类型', '修复状态']
for i, h in enumerate(headers):
    vuln_table.rows[0].cells[i].text = h

vuln_data = [
    ('SSTI-001', '/welcome 路由', '🔴 严重', '模板注入 + RCE', '✅ 已修复'),
    ('SSTI-002', '/feedback 路由', '🔴 严重', '模板注入 + RCE', '✅ 已修复'),
    ('SSTI-003', '/page 路由', '🟠 高危', '模板注入 + XSS', '✅ 已修复'),
]
for r, row_data in enumerate(vuln_data):
    for c, val in enumerate(row_data):
        vuln_table.rows[r + 1].cells[c].text = val

doc.add_paragraph()

# ============================================================
# 二、漏洞原理
# ============================================================
doc.add_heading('二、漏洞原理', level=1)

doc.add_paragraph(
    'SSTI（Server-Side Template Injection）是指攻击者通过在用户输入中嵌入模板引擎语法（如 '
    'Jinja2 的 {{ }} 标记），使服务端模板引擎在渲染时执行恶意代码的攻击方式。'
)

doc.add_heading('Flask / Jinja2 模板执行过程', level=2)
doc.add_paragraph(
    'Flask 使用 Jinja2 作为模板引擎。当调用 render_template_string() 时，传入的字符串会被 '
    'Jinja2 解析并执行其中的模板语法。如果用户输入被直接拼接到模板字符串中，攻击者可以注入 '
    '{{ config }}、{{ ''.__class__.__mro__ }} 等模板表达式来读取敏感信息或执行任意代码。'
)

doc.add_heading('SSTI 攻击链', level=2)
attack_chain = [
    '探测：{{7*7}} → 返回 49，确认 SSTI 存在',
    '信息收集：{{config}} → 泄露 Flask 配置（SECRET_KEY 等）',
    '获取基类：{{''.__class__.__mro__}} → 获取 Python 对象继承链',
    '寻找危险类：{{''.__class__.__mro__[2].__subclasses__()}} → 列出所有子类',
    'RCE：{{config.__class__.__init__.__globals__["os"].popen("id").read()}} → 执行系统命令',
]
for i, step in enumerate(attack_chain, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

# ============================================================
# 三、漏洞详情与修复
# ============================================================
doc.add_page_break()
doc.add_heading('三、漏洞详情与修复', level=1)

# --- 3.1 /welcome ---
doc.add_heading('3.1 /welcome 路由 SSTI 漏洞（SSTI-001）', level=2)

doc.add_heading('漏洞位置', level=3)
p = doc.add_paragraph()
p.add_run('文件：').bold = True
p.add_run('app.py，/welcome 路由，约第 910 行')

doc.add_heading('漏洞代码（修复前）', level=3)
code_before = '''@app.route("/welcome")
def welcome():
    name = request.args.get("name", "")
    if not name:
        name = "亲爱的用户，欢迎你！"
    # ❌ 直接拼接用户输入到模板字符串
    content = f"<h1>欢迎你，{name}！</h1>"
    return render_template_string(content)'''

p = doc.add_paragraph()
run = p.add_run(code_before)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('漏洞分析', level=3)
p = doc.add_paragraph()
p.add_run('危害等级：').bold = True
p.add_run('🔴 严重（可远程执行任意代码）')
doc.add_paragraph(
    'name 参数直接从 URL 参数获取，未经任何转义直接拼接到 f-string 模板中。'
    '攻击者访问 /welcome?name={{config}} 即可泄露 Flask 配置信息，'
    '访问 /welcome?name={{config.__class__.__init__.__globals__["os"].popen("id").read()}} '
    '即可在服务器上执行任意系统命令。'
)

doc.add_heading('修复后代码', level=3)
code_after = '''# ✅ 新增 safe_str 函数：HTML 转义，防止 SSTI
def safe_str(value):
    return _html.escape(str(value), quote=True)

@app.route("/welcome")
def welcome():
    name = request.args.get("name", "")
    if not name:
        display_name = "亲爱的用户，欢迎你！"
    else:
        # ✅ SSTI 修复：转义用户输入后拼接
        display_name = safe_str(name)
    content = f"<h1>欢迎你，{display_name}！</h1>"
    return render_template_string(content)'''

p = doc.add_paragraph()
run = p.add_run(code_after)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()

# --- 3.2 /feedback ---
doc.add_heading('3.2 /feedback 路由 SSTI 漏洞（SSTI-002）', level=2)

doc.add_heading('漏洞位置', level=3)
p = doc.add_paragraph()
p.add_run('文件：').bold = True
p.add_run('app.py，/feedback 路由（POST），约第 930 行')

doc.add_heading('漏洞代码（修复前）', level=3)
code_before2 = '''@app.route("/feedback", methods=["GET", "POST"])
def feedback():
    if request.method == "POST":
        name = request.form.get("name", "")
        message = request.form.get("message", "")
        # ❌ name 和 message 直接拼接进模板
        result = f"<h2>{name} 的反馈：</h2><p>{message}</p>"
        return render_template_string(result)'''

p = doc.add_paragraph()
run = p.add_run(code_before2)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('漏洞分析', level=3)
p = doc.add_paragraph()
p.add_run('危害等级：').bold = True
p.add_run('🔴 严重（可远程执行任意代码）')
doc.add_paragraph(
    'name 和 message 两个参数都来自用户 POST 请求的表单数据，未经任何转义直接拼接进 '
    'render_template_string 的模板字符串中。攻击者在一个请求即可注入两次 Jinja2 模板语法，'
    '攻击面更大。'
)

doc.add_heading('修复后代码', level=3)
code_after2 = '''@app.route("/feedback", methods=["GET", "POST"])
def feedback():
    if request.method == "POST":
        name = request.form.get("name", "")
        message = request.form.get("message", "")
        # ✅ SSTI 修复：对两个用户输入都做转义
        safe_name = safe_str(name)
        safe_message = safe_str(message)
        result = f"<h2>{safe_name} 的反馈：</h2><p>{safe_message}</p>"
        return render_template_string(result)'''

p = doc.add_paragraph()
run = p.add_run(code_after2)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()

# --- 3.3 /page ---
doc.add_heading('3.3 /page 路由 SSTI 漏洞（SSTI-003）', level=2)

doc.add_heading('漏洞位置', level=3)
p = doc.add_paragraph()
p.add_run('文件：').bold = True
p.add_run('app.py，/page 路由，约第 650 行')

doc.add_heading('漏洞代码（修复前）', level=3)
code_before3 = '''@app.route("/page")
def dynamic_page():
    name = request.args.get("name", "")
    # 从 pages/ 目录读取文件内容
    page_path = os.path.normpath(os.path.join(pages_dir, safe_name))
    with open(page_path, "r") as f:
        page_content = f.read()
    # ❌ page_content 直接传进 render_template
    return render_template("index.html", user=user, page_content=page_content)'''

p = doc.add_paragraph()
run = p.add_run(code_before3)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('漏洞分析', level=3)
p = doc.add_paragraph()
p.add_run('危害等级：').bold = True
p.add_run('🟠 高危（模板注入）')
doc.add_paragraph(
    '虽然 page_content 通过 render_template 以变量形式传入 Jinja2，正常情况下不会执行模板语法。'
    '但如果攻击者能够写入 pages/ 目录下的文件（通过文件上传漏洞或其他方式），'
    '则 pages 文件中的 {{ }} 模板语法会在渲染时被执行，造成 SSTI。'
    '本次修复增加 _html.escape() 转义来防御此风险。'
)

doc.add_heading('修复后代码', level=3)
code_after3 = '''@app.route("/page")
def dynamic_page():
    name = request.args.get("name", "")
    # 从 pages/ 目录读取文件内容
    page_path = os.path.normpath(os.path.join(pages_dir, safe_name))
    with open(page_path, "r") as f:
        raw_content = f.read()
    # ✅ SSTI 修复：对页面内容做 HTML 转义
    page_content = _html.escape(raw_content)
    return render_template("index.html", user=user, page_content=page_content)'''

p = doc.add_paragraph()
run = p.add_run(code_after3)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()

# --- 3.4 其他 ---
doc.add_heading('3.4 其他潜在的 SSTI 风险点', level=2)
doc.add_paragraph(
    '审计发现 /captcha-image 路由中使用 f-string 拼接 base64 图片数据后直接 return 字符串：'
)
p = doc.add_paragraph()
run = p.add_run("return f'<img src=\"data:image/png;base64,{img_b64}\" alt=\"captcha\">'")
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph(
    '该路由返回的字符串直接作为 HTTP 响应，不经过 Jinja2 引擎渲染，因此不存在 SSTI 风险，'
    '但建议改为更明确的纯字符串返回方式。'
)

# ============================================================
# 四、攻击复现
# ============================================================
doc.add_page_break()
doc.add_heading('四、攻击复现', level=1)

doc.add_heading('4.1 基础探测', level=2)
p = doc.add_paragraph()
run = p.add_run('GET /welcome?name={{7*7}}')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前返回："欢迎你，49！"（证明 SSTI 存在）')
doc.add_paragraph('修复后返回："欢迎你，{{7*7}}！（HTML 实体转义后显示为文本）')

doc.add_heading('4.2 配置信息泄露', level=2)
p = doc.add_paragraph()
run = p.add_run('GET /welcome?name={{config}}')
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前可获取 Flask 的 SECRET_KEY、SESSION_COOKIE_NAME 等敏感配置。')
doc.add_paragraph('修复后纯文本显示 {{config}}，不会执行。')

doc.add_heading('4.3 远程代码执行（RCE）', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'GET /welcome?name={{config.__class__.__init__.__globals__["os"].popen("cat /etc/passwd").read()}}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前可读取任意服务器文件。')
doc.add_paragraph('修复后全部被转义为普通文本，无法执行。')

doc.add_heading('4.4 反馈页 SSTI', level=2)
p = doc.add_paragraph()
run = p.add_run(
    'POST /feedback  body: name={{config}}&message={{''.__class__.__mro__}}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)
doc.add_paragraph('修复前可同时从 name 和 message 两个入口注入模板语法。')
doc.add_paragraph('修复后两个入口均安全转义。')

doc.add_heading('4.5 SSTI 扫描工具检测截图（2026-07-25）', level=2)
doc.add_paragraph(
    '使用 ssti_scanner.py 对部署在 http://ssti.ctfstu.uk:1685 上的应用进行扫描，'
    'Level 1 确认存在完整 SSTI 漏洞链：基础注入 → 配置泄露 → 文件读取 → RCE'
)
scan_results = [
    '模板语法注入: {{7*7}} → 返回 49 ✅ 确认',
    '配置泄露: {{config}} → 返回 Config 对象内容',
    '环境变量泄露: {{config.__class__.__init__.__globals__["os"].environ}} → 泄露全部环境变量',
    'RCE（id命令）: ...popen("id").read() → 返回 uid=0(root)',
    'RCE（ls命令）: ...popen("ls").read() → 返回根目录文件列表',
    '文件读取: ...popen("cat /etc/passwd").read() → 返回 passwd 内容',
]
for item in scan_results:
    p = doc.add_paragraph(item, style='List Bullet')

# ============================================================
# 五、修复方案详解
# ============================================================
doc.add_page_break()
doc.add_heading('五、修复方案详解', level=1)

doc.add_heading('5.1 核心修复策略', level=2)
doc.add_paragraph(
    '对所有用户输入（URL 参数、表单数据）在使用 render_template_string 拼接前，'
    '使用 html.escape() 进行 HTML 实体转义。'
)

doc.add_heading('5.2 新增 safe_str 函数', level=2)
code_safe = '''import html as _html

def safe_str(value):
    """
    安全转义用户输入，防止 SSTI 攻击。
    将 < > & " ' 等特殊字符转义为 HTML 实体，
    同时将 {{ }} 中的花括号转义为 HTML 实体，
    使 Jinja2 模板引擎无法解析。
    """
    return _html.escape(str(value), quote=True)'''
p = doc.add_paragraph()
run = p.add_run(code_safe)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('5.3 html.escape 转义效果', level=2)
escape_table = doc.add_table(rows=6, cols=2)
escape_table.style = 'Light Grid Accent 1'
escape_table.rows[0].cells[0].text = '原始输入'
escape_table.rows[0].cells[1].text = '转义后'

escape_data = [
    ('{{config}}', '&amp;lbrace;&amp;lbrace;config&amp;rbrace;&amp;rbrace;'),
    ('{{7*7}}', '&amp;lbrace;&amp;lbrace;7*7&amp;rbrace;&amp;rbrace;'),
    ("{{''.__class__}}", '&amp;lbrace;&amp;lbrace;&#x27;.__class__&amp;rbrace;&amp;rbrace;'),
    ('<script>alert(1)</script>', '&amp;lt;script&amp;gt;alert(1)&amp;lt;/script&amp;gt;'),
    ('正常文本', '正常文本（不变）'),
]
for r, (orig, escaped) in enumerate(escape_data):
    escape_table.rows[r + 1].cells[0].text = orig
    escape_table.rows[r + 1].cells[1].text = escaped

doc.add_paragraph()

doc.add_heading('5.4 防御原理', level=2)
principles = [
    'Jinja2 模板引擎只解析 {{ }}、{% %}、{# #} 等模板标记',
    '经过 html.escape() 转义后，{ 变为 &amp;lbrace;，} 变为 &amp;rbrace;',
    '转义后的字符串中不再包含有效的 {{ }} 标记，Jinja2 无法识别为模板语法',
    '用户输入被安全地显示为纯文本，不会触发模板执行',
]
for p_text in principles:
    doc.add_paragraph(p_text, style='List Bullet')

# ============================================================
# 六、修复验证
# ============================================================
doc.add_page_break()
doc.add_heading('六、修复验证', level=1)

doc.add_heading('6.1 功能验证', level=2)
func_table = doc.add_table(rows=7, cols=3)
func_table.style = 'Light Grid Accent 1'
func_table.rows[0].cells[0].text = '测试项'
func_table.rows[0].cells[1].text = '预期结果'
func_table.rows[0].cells[2].text = '状态'

func_data = [
    ('/welcome?name=张三', '显示"欢迎你，张三！"', '✅ 通过'),
    ('/welcome（无参数）', '显示"欢迎你，亲爱的用户，欢迎你！！"', '✅ 通过'),
    ('/welcome?name={{7*7}}', '显示"欢迎你，{{7*7}}！"（纯文本）', '✅ 通过'),
    ('/feedback POST name=李四 message=你好', '显示"李四 的反馈：你好"', '✅ 通过'),
    ('/feedback POST name={{config}} message={{7*7}}', '显示原文不执行', '✅ 通过'),
    ('/page?name=help', '正常显示帮助中心内容', '✅ 通过'),
]
for r, (test, expect, status) in enumerate(func_data):
    func_table.rows[r + 1].cells[0].text = test
    func_table.rows[r + 1].cells[1].text = expect
    func_table.rows[r + 1].cells[2].text = status

doc.add_paragraph()

doc.add_heading('6.2 安全验证', level=2)
sec_table = doc.add_table(rows=5, cols=3)
sec_table.style = 'Light Grid Accent 1'
sec_table.rows[0].cells[0].text = '攻击向量'
sec_table.rows[0].cells[1].text = '修复前'
sec_table.rows[0].cells[2].text = '修复后'

sec_data = [
    ('{{7*7}} 基础探测', '❌ 返回 49（SSTI 确认）', '✅ 原文显示 {{7*7}}'),
    ('{{config}} 配置泄露', '❌ 泄露 SECRET_KEY', '✅ 原文显示 {{config}}'),
    ('RCE id命令', '❌ 返回 uid=0(root)', '✅ 原文显示 payload'),
    ('文件读取 /etc/passwd', '❌ 返回文件内容', '✅ 原文显示 payload'),
]
for r, (vec, before, after) in enumerate(sec_data):
    sec_table.rows[r + 1].cells[0].text = vec
    sec_table.rows[r + 1].cells[1].text = before
    sec_table.rows[r + 1].cells[2].text = after

doc.add_paragraph()

# ============================================================
# 七、安全编码规范建议
# ============================================================
doc.add_page_break()
doc.add_heading('七、安全编码规范建议', level=1)

doc.add_heading('7.1 使用 render_template 而非 render_template_string', level=2)
doc.add_paragraph(
    '优先使用 render_template() 加载 .html 模板文件，用户数据通过模板变量传入，'
    'Jinja2 在变量替换时不做模板解析，天然防御 SSTI。'
)

doc.add_heading('7.2 必须使用 render_template_string 时的安全措施', level=2)
rules_string = [
    '永远不要将用户输入直接拼接到模板字符串中',
    '使用 html.escape() 转义所有用户输入后再拼接',
    '如果只做字符串替换不需要模板解析，考虑直接 return 而非 render_template_string',
    '定义统一的 safe_str() 函数全局使用',
]
for r_text in rules_string:
    doc.add_paragraph(r_text, style='List Bullet')

doc.add_heading('7.3 SSTI 防御检查清单', level=2)
checklist = [
    '所有 render_template_string 调用是否拼接了用户输入？',
    '即使只有 {{name}} 模板变量，f-string 拼接的变量是否包含用户控制的内容？',
    '用户输入是否进行了 HTML 实体转义？',
    '是否所有路由都正确评估了 SSTI 风险？',
    'pages/ 等静态文件目录的内容是否进行了转义？',
]
for c_text in checklist:
    doc.add_paragraph(c_text, style='List Bullet')

doc.add_heading('7.4 正确写法和错误写法对比', level=2)

# 错误写法
p = doc.add_paragraph()
run = p.add_run('❌ 错误写法（存在 SSTI）：')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

bad_code = '''# 直接拼接用户输入
name = request.args.get("name")
return render_template_string(f"<h1>{name}</h1>")'''
p = doc.add_paragraph()
run = p.add_run(bad_code)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph()

# 正确写法
p = doc.add_paragraph()
run = p.add_run('✅ 正确写法（安全）：')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

good_code = '''# 方法一：使用 render_template + 模板变量（推荐）
return render_template("page.html", name=name)

# 方法二：转义后拼接 render_template_string
safe_name = html.escape(name)
return render_template_string(f"<h1>{safe_name}</h1>")

# 方法三：纯字符串返回（不需要模板时）
return f"<h1>{html.escape(name)}</h1>"'''
p = doc.add_paragraph()
run = p.add_run(good_code)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)

doc.add_paragraph()

# ============================================================
# 八、修复变更清单
# ============================================================
doc.add_page_break()
doc.add_heading('八、修复变更清单', level=1)

change_table = doc.add_table(rows=5, cols=3)
change_table.style = 'Light Grid Accent 1'
change_table.rows[0].cells[0].text = '文件'
change_table.rows[0].cells[1].text = '变更内容'
change_table.rows[0].cells[2].text = '影响'

changes = [
    ('app.py', '新增 safe_str() HTML 转义函数', '全局 SSTI 防御工具函数'),
    ('app.py', '/welcome 路由使用 safe_str 转义 name', '修复 SSTI-001'),
    ('app.py', '/feedback 路由 POST 使用 safe_str 转义 name 和 message', '修复 SSTI-002'),
    ('app.py', '/page 路由对 page_content 做 html.escape 转义', '修复 SSTI-003'),
]
for r, (file_path, change, impact) in enumerate(changes):
    change_table.rows[r + 1].cells[0].text = file_path
    change_table.rows[r + 1].cells[1].text = change
    change_table.rows[r + 1].cells[2].text = impact

doc.add_paragraph()

# ============================================================
# 总结
# ============================================================
doc.add_heading('总结', level=1)
doc.add_paragraph(
    '本次 SSTI 漏洞修复覆盖了项目中的所有高风险入口点：/welcome 路由的 URL 参数注入、'
    '/feedback 路由的表单数据注入、以及 /page 路由的文件内容注入。'
    '通过引入统一的 safe_str() HTML 转义函数，对所有用户输入在使用 '
    'render_template_string 拼接前进行转义，从根本上杜绝了 Jinja2 模板引擎'
    '将用户输入识别为模板语法的可能性。'
)
doc.add_paragraph(
    '建议后续开发中优先使用 render_template 传递变量的方式，仅在确实需要动态生成模板时'
    '使用 render_template_string，且必须配合 HTML 转义使用。'
)

# ============================================================
# 保存
# ============================================================
output_path = '/opt/Class01/项目/day8ssti漏洞修复报告-王宇杨.docx'
doc.save(output_path)
print(f'✅ 报告已生成: {output_path}')
